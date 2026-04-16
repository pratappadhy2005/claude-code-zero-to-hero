"""
JAC MEGA Test — Grade 5 General Ability Practice Worksheet Generator
Clean, readable layout using python-docx best practices.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def hex_to_rgb(hex_str: str) -> RGBColor:
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def shade_cell(cell, fill_hex: str):
    """Apply solid background colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shd elements
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def hide_borders(table):
    """Remove all borders from a table."""
    tbl_pr = table._tbl.tblPr
    # Remove existing tblBorders
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:sz"), "0")
        tb.append(el)
    tbl_pr.append(tb)


def set_col_width(table, col_idx: int, width_cm: float):
    """Set explicit width on every cell in a column."""
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:tcW")):
            tc_pr.remove(old)
        tcw = OxmlElement("w:tcW")
        # Convert cm to twentieths of a point (twips): 1 cm = 567 twips
        twips = int(width_cm * 567)
        tcw.set(qn("w:w"), str(twips))
        tcw.set(qn("w:type"), "dxa")
        tc_pr.append(tcw)


def cell_vert_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tc_pr.append(va)


def para_space(para, before_pt: float = 0, after_pt: float = 0, line_pt: float = None):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"),  str(int(after_pt * 20)))
    if line_pt:
        spc.set(qn("w:line"), str(int(line_pt * 20)))
        spc.set(qn("w:lineRule"), "exact")
    pPr.append(spc)


def rn(para, text, bold=False, size=11, colour=None, italic=False):
    """Add a run to paragraph."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = hex_to_rgb(colour)
    return run


# ── Colours ──────────────────────────────────────────────────────────────────
RED    = "C0392B"
AMBER  = "E67E22"
DARK   = "1A5276"
LIGHT  = "EAF4FF"
RED_L  = "FADBD8"
AMB_L  = "FDEBD0"
ALT    = "EBF5FB"
WHITE  = "FFFFFF"
GREY_L = "F2F3F4"


# ── Section header ────────────────────────────────────────────────────────────
def section_header(doc, title: str, tag: str, fill: str):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hide_borders(tbl)
    set_col_width(tbl, 0, 13.5)
    set_col_width(tbl, 1, 3.5)

    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    shade_cell(lc, fill)
    shade_cell(rc, fill)
    cell_vert_center(lc)
    cell_vert_center(rc)

    # Padding via paragraph indent
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rn(lp, f"  {title}", bold=True, size=12, colour=WHITE)
    para_space(lp, before_pt=5, after_pt=5)

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn(rp, f"{tag}  ", bold=True, size=9, colour=WHITE)
    para_space(rp, before_pt=5, after_pt=5)

    gap = doc.add_paragraph()
    para_space(gap, after_pt=2)


# ── Topic note ────────────────────────────────────────────────────────────────
def topic_note(doc, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    hide_borders(tbl)
    cell = tbl.cell(0, 0)
    shade_cell(cell, LIGHT)
    p = cell.paragraphs[0]
    rn(p, text, size=9.5, italic=True, colour="2C3E50")
    para_space(p, before_pt=3, after_pt=3)
    gap = doc.add_paragraph()
    para_space(gap, after_pt=2)


# ── MCQ block ─────────────────────────────────────────────────────────────────
def mcq_section(doc, questions: list[dict], start_num: int = 1):
    """
    Each question uses a 2-row, 2-col table:
      Row 0: Q-number | Question text
      Row 1: empty   | A) ... B) ... C) ... D) ...
    This keeps question and options visually glued together.
    """
    for i, item in enumerate(questions):
        q_num = start_num + i

        tbl = doc.add_table(rows=2, cols=2)
        hide_borders(tbl)
        set_col_width(tbl, 0, 1.0)
        set_col_width(tbl, 1, 16.0)

        # Alternating light background
        bg = GREY_L if i % 2 == 0 else WHITE
        for r in tbl.rows:
            for c in r.cells:
                shade_cell(c, bg)

        # Row 0: number + question
        num_cell = tbl.cell(0, 0)
        qtext_cell = tbl.cell(0, 1)

        np = num_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn(np, f"{q_num}.", bold=True, size=11, colour="1A1A2E")
        para_space(np, before_pt=5, after_pt=1)

        qp = qtext_cell.paragraphs[0]
        qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rn(qp, f"  {item['q']}", size=11, colour="1A1A2E")
        para_space(qp, before_pt=5, after_pt=1)

        # Row 1: options
        opt_cell = tbl.cell(1, 1)
        op = opt_cell.paragraphs[0]
        op.alignment = WD_ALIGN_PARAGRAPH.LEFT
        labels = ["A", "B", "C", "D"]
        for j, opt in enumerate(item["opts"]):
            rn(op, f"  ({labels[j]}) {opt}    ", size=10.5, colour="2C3E50")
        para_space(op, before_pt=1, after_pt=5)

        spacer = doc.add_paragraph()
        para_space(spacer, after_pt=1)


# ── Answer key ────────────────────────────────────────────────────────────────
def answer_key(doc, sections: list[dict]):
    doc.add_page_break()

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p, "ANSWER KEY", bold=True, size=16, colour=DARK)
    para_space(p, after_pt=2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p2, "For Parent / Teacher Use Only", italic=True, size=10, colour="555555")
    para_space(p2, after_pt=8)

    # Flatten all answers with section labels
    rows_data = []
    for sec in sections:
        for q_short, ans, expl in sec["answers"]:
            rows_data.append((sec["topic"], q_short, ans, expl))

    # Table
    tbl = doc.add_table(rows=1 + len(rows_data), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    col_widths = [2.8, 6.5, 1.4, 6.3]
    for ci, w in enumerate(col_widths):
        set_col_width(tbl, ci, w)

    # Header row
    headers = ["Topic", "Question (short)", "Ans", "Explanation"]
    for ci, hdr in enumerate(headers):
        cell = tbl.cell(0, ci)
        shade_cell(cell, DARK)
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(hp, hdr, bold=True, size=9.5, colour=WHITE)
        para_space(hp, before_pt=3, after_pt=3)

    # Data rows
    for ri, (topic, q_short, ans, expl) in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        bg = ALT if ri % 2 == 0 else WHITE
        vals = [topic, q_short, ans, expl]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            cp = cell.paragraphs[0]
            cp.alignment = aligns[ci]
            bold = ci == 2  # bold the answer letter
            colour = RED if (ci == 2) else "1A1A2E"
            rn(cp, val, bold=bold, size=9, colour=colour)
            para_space(cp, before_pt=2, after_pt=2)


# ── Scoring guide ─────────────────────────────────────────────────────────────
def scoring_guide(doc):
    doc.add_paragraph()
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(hdr, "SCORING GUIDE", bold=True, size=13, colour=DARK)
    para_space(hdr, before_pt=8, after_pt=6)

    bands = [
        ("30–36 / 36  (83%+)",  "Excellent! Focus on speed. Try to complete each section in under 5 minutes."),
        ("24–29 / 36  (67–80%)", "Good progress. Review every wrong answer using the key above. Keep practising daily."),
        ("18–23 / 36  (50–66%)", "You are building your vocabulary. Re-read each section note, then redo the questions."),
        ("Below 18  (<50%)",     "Focus on Sections 1–4 first (Priority topics). Redo each section before moving on."),
    ]

    tbl = doc.add_table(rows=1 + len(bands), cols=2)
    tbl.style = "Table Grid"
    set_col_width(tbl, 0, 5.5)
    set_col_width(tbl, 1, 11.5)

    # Header
    for ci, hd in enumerate(["Score Band", "What To Do"]):
        cell = tbl.cell(0, ci)
        shade_cell(cell, DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, hd, bold=True, size=10, colour=WHITE)
        para_space(p, before_pt=3, after_pt=3)

    for ri, (band, action) in enumerate(bands):
        bg = ALT if ri % 2 == 0 else WHITE
        row = tbl.rows[ri + 1]
        shade_cell(row.cells[0], bg)
        shade_cell(row.cells[1], bg)

        p0 = row.cells[0].paragraphs[0]
        rn(p0, band, bold=True, size=9.5, colour=DARK)
        para_space(p0, before_pt=3, after_pt=3)

        p1 = row.cells[1].paragraphs[0]
        rn(p1, action, size=9.5, colour="1A1A2E")
        para_space(p1, before_pt=3, after_pt=3)


# ── QUESTION DATA ─────────────────────────────────────────────────────────────

SCIENCE = dict(
    questions=[
        {"q": "Which process do plants use to make their own food using sunlight?",
         "opts": ["Respiration", "Photosynthesis", "Digestion", "Germination"]},
        {"q": "What is the name for animals that eat only plants?",
         "opts": ["Carnivores", "Omnivores", "Herbivores", "Scavengers"]},
        {"q": "Which part of a flower produces pollen?",
         "opts": ["Petal", "Pistil", "Stamen", "Sepal"]},
        {"q": "An animal that has a backbone is called a ___.",
         "opts": ["Invertebrate", "Vertebrate", "Mammal", "Reptile"]},
        {"q": "Which of the following is NOT a renewable energy source?",
         "opts": ["Solar", "Wind", "Coal", "Hydroelectric"]},
        {"q": "The layer of gases surrounding the Earth is called the ___.",
         "opts": ["Hydrosphere", "Lithosphere", "Atmosphere", "Biosphere"]},
        {"q": "What do we call the change from a liquid to a gas?",
         "opts": ["Condensation", "Evaporation", "Melting", "Freezing"]},
        {"q": "Which organ pumps blood around the human body?",
         "opts": ["Lungs", "Brain", "Heart", "Liver"]},
    ],
    answers=[
        ("Plants make food using sunlight?", "B", "Photosynthesis uses sunlight + water + CO₂ to make glucose and oxygen."),
        ("Animals that eat only plants?", "C", "Herbivores eat only plants. Carnivores eat meat; omnivores eat both."),
        ("Part of flower that produces pollen?", "C", "The stamen (male part) produces pollen. Pistil is the female part."),
        ("Animal with a backbone?", "B", "Vertebrates have a backbone. Invertebrates (e.g. insects) do not."),
        ("NOT a renewable energy source?", "C", "Coal is a fossil fuel — non-renewable. Solar, wind, and hydro are renewable."),
        ("Layer of gases around Earth?", "C", "The atmosphere is the layer of gases surrounding Earth."),
        ("Change from liquid to gas?", "B", "Evaporation = liquid → gas. Condensation is the reverse."),
        ("Organ that pumps blood?", "C", "The heart pumps blood throughout the body via blood vessels."),
    ]
)

SYNONYMS = dict(
    questions=[
        {"q": "Which word is closest in meaning to ANCIENT?",
         "opts": ["New", "Old", "Broken", "Slow"]},
        {"q": "Which word is closest in meaning to COURAGEOUS?",
         "opts": ["Careful", "Selfish", "Brave", "Proud"]},
        {"q": "Which word is closest in meaning to RAPID?",
         "opts": ["Loud", "Quick", "Huge", "Sharp"]},
        {"q": "Which word means the same as PECULIAR?",
         "opts": ["Common", "Strange", "Bright", "Quiet"]},
        {"q": "Which word is closest in meaning to WEARY?",
         "opts": ["Happy", "Hungry", "Tired", "Strong"]},
        {"q": "Which word means the same as COMMENCE?",
         "opts": ["Finish", "Begin", "Break", "Sleep"]},
    ],
    answers=[
        ("Closest to ANCIENT?", "B", "Old = ancient. New is the antonym."),
        ("Closest to COURAGEOUS?", "C", "Brave = courageous. Careful and proud are different."),
        ("Closest to RAPID?", "B", "Quick = rapid. Rapid means fast."),
        ("Same as PECULIAR?", "B", "Strange = peculiar (unusual, odd)."),
        ("Closest to WEARY?", "C", "Tired = weary. Weary means exhausted."),
        ("Same as COMMENCE?", "B", "Begin = commence. Commence means to start."),
    ]
)

ANALOGIES = dict(
    questions=[
        {"q": "Puppy is to dog as kitten is to ___.",
         "opts": ["Cat", "Rabbit", "Lion", "Fish"]},
        {"q": "Gloves are to hands as socks are to ___.",
         "opts": ["Arms", "Feet", "Head", "Ears"]},
        {"q": "Painter is to brush as writer is to ___.",
         "opts": ["Canvas", "Story", "Pen", "Ink"]},
        {"q": "Wet is to dry as cold is to ___.",
         "opts": ["Icy", "Warm", "Windy", "Dark"]},
        {"q": "Hospital is to doctor as school is to ___.",
         "opts": ["Student", "Book", "Teacher", "Lesson"]},
        {"q": "Bark is to tree as shell is to ___.",
         "opts": ["Sea", "Stone", "Egg", "Rock"]},
    ],
    answers=[
        ("Puppy : dog :: kitten : ___?", "A", "A puppy grows into a dog; a kitten grows into a cat."),
        ("Gloves : hands :: socks : ___?", "B", "Gloves cover hands; socks cover feet."),
        ("Painter : brush :: writer : ___?", "C", "A painter uses a brush; a writer uses a pen."),
        ("Wet : dry :: cold : ___?", "B", "Wet and dry are opposites; cold and warm are opposites."),
        ("Hospital : doctor :: school : ___?", "C", "A doctor works in a hospital; a teacher works in a school."),
        ("Bark : tree :: shell : ___?", "C", "Bark covers a tree; shell covers an egg."),
    ]
)

PALINDROMES = dict(
    questions=[
        {"q": "A palindrome reads the same forwards AND backwards. Which word is a palindrome?",
         "opts": ["Level", "Apple", "Bridge", "Stone"]},
        {"q": "Which of the following is a palindrome?",
         "opts": ["River", "Noon", "Clock", "Ocean"]},
        {"q": "Which word is NOT a palindrome?",
         "opts": ["Racecar", "Civic", "Radar", "Planet"]},
        {"q": "Which number is a palindrome?",
         "opts": ["123", "456", "121", "789"]},
        {"q": "How many letters does the palindrome 'MADAM' have?",
         "opts": ["4", "5", "6", "7"]},
        {"q": "Which word is a palindrome?",
         "opts": ["Garden", "Kayak", "Flower", "Window"]},
    ],
    answers=[
        ("Which word is a palindrome?", "A", "L-E-V-E-L — reads the same forwards and backwards."),
        ("Which is a palindrome?", "B", "N-O-O-N — a classic palindrome."),
        ("Which is NOT a palindrome?", "D", "P-L-A-N-E-T → backwards is T-E-N-A-L-P. Not a palindrome."),
        ("Which number is a palindrome?", "C", "121 → one-two-one, the same both ways."),
        ("How many letters in MADAM?", "B", "M-A-D-A-M has 5 letters and is a palindrome."),
        ("Which word is a palindrome?", "B", "K-A-Y-A-K — reads the same forwards and backwards."),
    ]
)

IDIOMS = dict(
    questions=[
        {"q": "'It's raining cats and dogs' means ___.",
         "opts": ["Animals are falling", "It is raining heavily", "The weather is cold", "A storm is coming"]},
        {"q": "If someone is 'barking up the wrong tree', they are ___.",
         "opts": ["Talking to animals", "Looking in the wrong place", "Being noisy", "Climbing a tree"]},
        {"q": "'Hit the sack' means to ___.",
         "opts": ["Go to sleep", "Hit a bag", "Start running", "Eat dinner"]},
        {"q": "'The ball is in your court' means ___.",
         "opts": ["Play sport now", "It is your turn to decide", "You are on a tennis court", "The game is over"]},
        {"q": "'Under the weather' means ___.",
         "opts": ["Standing in the rain", "Feeling unwell", "Outside in a storm", "Very cold"]},
        {"q": "'Break a leg' is said to someone who is ___.",
         "opts": ["Injured", "About to perform", "Running a race", "Going to hospital"]},
    ],
    answers=[
        ("Raining cats and dogs?", "B", "This idiom means raining very heavily — nothing to do with animals."),
        ("Barking up the wrong tree?", "B", "Means pursuing the wrong course of action or the wrong person."),
        ("Hit the sack?", "A", "To 'hit the sack' means to go to bed / go to sleep."),
        ("Ball is in your court?", "B", "From tennis — the next decision or action belongs to you."),
        ("Under the weather?", "B", "Feeling unwell or slightly ill."),
        ("Break a leg?", "B", "Said to performers before they go on stage — it means good luck."),
    ]
)

COLLOQUIALISMS = dict(
    questions=[
        {"q": "In Australian English, 'arvo' means ___.",
         "opts": ["Morning", "Afternoon", "Evening", "Weekend"]},
        {"q": "'Fortnight' means ___.",
         "opts": ["One week", "Two weeks", "One month", "Ten days"]},
        {"q": "If something is described as 'dodgy', it means it is ___.",
         "opts": ["Excellent", "Suspicious or poor quality", "Very fast", "Very heavy"]},
        {"q": "In British English, a 'lift' is what Australians call an ___.",
         "opts": ["Elevator", "Escalator", "Aeroplane", "Exit"]},
        {"q": "'Reckon' is used in Australian English to mean ___.",
         "opts": ["Count money", "Think or believe", "Run quickly", "Ask a question"]},
        {"q": "In Australian English, 'brekkie' refers to ___.",
         "opts": ["A break from school", "Breakfast", "A bicycle", "A biscuit"]},
    ],
    answers=[
        ("Arvo?", "B", "'Arvo' = afternoon in informal Australian English."),
        ("Fortnight?", "B", "'Fortnight' = two weeks (14 nights). British and Australian usage."),
        ("Dodgy?", "B", "'Dodgy' = suspicious, unreliable, or poor quality."),
        ("British 'lift'?", "A", "British 'lift' = Australian/American 'elevator'."),
        ("Reckon?", "B", "'Reckon' = to think or believe in informal Australian English."),
        ("Brekkie?", "B", "'Brekkie' = breakfast in informal Australian English."),
    ]
)


# ── BUILD ─────────────────────────────────────────────────────────────────────

def build(output_path: str):
    doc = Document()

    # A4 page setup
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin  = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin   = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Default paragraph style: no extra spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── COVER ──────────────────────────────────────────────────────────────

    spacer = doc.add_paragraph()
    para_space(spacer, after_pt=16)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(t, "GENERAL ABILITY", bold=True, size=22, colour="1A1A2E")
    para_space(t, after_pt=4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(t2, "PRACTICE WORKSHEET", bold=True, size=18, colour="1A1A2E")
    para_space(t2, after_pt=8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(sub, "James An College  MEGA Test  |  Grade 5  |  General Ability", size=11, colour="555555")
    para_space(sub, after_pt=16)

    # Student fields table
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_width(tbl, 0, 4.0)
    set_col_width(tbl, 1, 4.0)
    set_col_width(tbl, 2, 4.0)
    set_col_width(tbl, 3, 4.0)
    labels_row = ["Student Name", "Date", "Score  / 36", "Time Taken"]
    for ci, lbl in enumerate(labels_row):
        shade_cell(tbl.cell(0, ci), DARK)
        p = tbl.cell(0, ci).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, lbl, bold=True, size=10, colour=WHITE)
        para_space(p, before_pt=3, after_pt=3)

        shade_cell(tbl.cell(1, ci), LIGHT)
        p2 = tbl.cell(1, ci).paragraphs[0]
        rn(p2, " ", size=14)  # blank row for writing
        para_space(p2, before_pt=6, after_pt=6)

    spacer2 = doc.add_paragraph()
    para_space(spacer2, after_pt=10)

    # Instructions box
    itbl = doc.add_table(rows=1, cols=1)
    hide_borders(itbl)
    ic = itbl.cell(0, 0)
    shade_cell(ic, LIGHT)
    ip = ic.paragraphs[0]
    rn(ip, "Instructions:  ", bold=True, size=11, colour=DARK)
    rn(ip, "Circle the letter of the best answer. "
           "Work through each section carefully — they are ordered from most important to fix first. "
           "Check your answers using the Answer Key on the final page. "
           "Use Australian English spelling throughout.",
       size=10.5, colour="1A1A2E")
    para_space(ip, before_pt=5, after_pt=5)

    spacer3 = doc.add_paragraph()
    para_space(spacer3, after_pt=10)

    # Topics overview table
    ov = doc.add_table(rows=2, cols=6)
    hide_borders(ov)
    topics_row = ["Science Vocab", "Synonyms", "Analogies", "Palindromes", "Idioms", "Colloquialisms"]
    prio_row   = ["Priority ★★★", "Priority ★★★", "Priority ★★★", "Priority ★★★", "Improve ★★", "Improve ★★"]
    fills_row  = [RED, RED, RED, RED, AMBER, AMBER]
    light_fills = [RED_L, RED_L, RED_L, RED_L, AMB_L, AMB_L]
    col_w = 2.8
    for ci in range(6):
        set_col_width(ov, ci, col_w)
        c0 = ov.cell(0, ci)
        c1 = ov.cell(1, ci)
        shade_cell(c0, fills_row[ci])
        shade_cell(c1, light_fills[ci])

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p0, topics_row[ci], bold=True, size=9, colour=WHITE)
        para_space(p0, before_pt=3, after_pt=2)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p1, prio_row[ci], size=8, colour="1A1A2E")
        para_space(p1, before_pt=2, after_pt=3)

    doc.add_page_break()

    # ── SECTIONS ───────────────────────────────────────────────────────────

    q_offset = 1  # running question counter

    # Section 1: Science
    section_header(doc, "Section 1  —  Science Vocabulary  (8 questions)", "PRIORITY ★★★", RED)
    topic_note(doc,
        "You missed 5/8 Science questions (Q6, Q7, Q15, Q28, Q30). "
        "Q6, Q15 and Q30 were answered correctly by 86–90% of the cohort — these are "
        "\"should-know\" questions. Focus on: biology terms, energy types, Earth science vocabulary.")
    mcq_section(doc, SCIENCE["questions"], q_offset)
    q_offset += len(SCIENCE["questions"])
    doc.add_page_break()

    # Section 2: Synonyms
    section_header(doc, "Section 2  —  Synonyms  (6 questions)", "PRIORITY ★★★", RED)
    topic_note(doc,
        "You missed both Synonym questions (Q4, Q27). "
        "A synonym is a word with the same or nearly the same meaning as another word. "
        "Tip: think about the full meaning of the target word, then find the closest match.")
    mcq_section(doc, SYNONYMS["questions"], q_offset)
    q_offset += len(SYNONYMS["questions"])
    doc.add_page_break()

    # Section 3: Analogies
    section_header(doc, "Section 3  —  Analogies  (6 questions)", "PRIORITY ★★★", RED)
    topic_note(doc,
        "You missed the Analogy question (Q5 — only 28% of the cohort got it right, so it was hard). "
        "An analogy shows how two pairs of words relate: \"X is to Y as ___ is to ___\". "
        "Step 1: identify the relationship (e.g. young animal → adult). "
        "Step 2: apply that same relationship to the new pair.")
    mcq_section(doc, ANALOGIES["questions"], q_offset)
    q_offset += len(ANALOGIES["questions"])
    doc.add_page_break()

    # Section 4: Palindromes
    section_header(doc, "Section 4  —  Palindromes  (6 questions)", "PRIORITY ★★★", RED)
    topic_note(doc,
        "You missed both Palindrome questions (Q23, Q24). Q24 was answered correctly by 82% of students — "
        "a fast fix. Rule: a palindrome reads the same forwards AND backwards. "
        "Examples: racecar, level, noon, madam, kayak, civic. "
        "Tip: write the word backwards and compare letter by letter.")
    mcq_section(doc, PALINDROMES["questions"], q_offset)
    q_offset += len(PALINDROMES["questions"])
    doc.add_page_break()

    # Section 5: Idioms
    section_header(doc, "Section 5  —  Idioms  (6 questions)", "IMPROVE ★★", AMBER)
    topic_note(doc,
        "You missed the Idiom question (Q11 — 77% of the cohort got it right). "
        "An idiom is a phrase whose meaning is different from the literal words — "
        "you cannot work them out from the words alone. They must be memorised. "
        "Keep an 'idiom journal' and learn 5 new expressions per week.")
    mcq_section(doc, IDIOMS["questions"], q_offset)
    q_offset += len(IDIOMS["questions"])
    doc.add_page_break()

    # Section 6: Colloquialisms
    section_header(doc, "Section 6  —  Colloquialisms  (6 questions)", "IMPROVE ★★", AMBER)
    topic_note(doc,
        "You missed Q21 (Colloquialisms). "
        "A colloquialism is an informal word used in everyday conversation. "
        "This test focuses on Australian and British informal expressions. "
        "Learn: arvo, brekkie, fortnight, reckon, dodgy, cheeky, ta, footy.")
    mcq_section(doc, COLLOQUIALISMS["questions"], q_offset)

    # ── ANSWER KEY ─────────────────────────────────────────────────────────
    all_sections = [
        {"topic": "Science",        "answers": SCIENCE["answers"]},
        {"topic": "Synonyms",       "answers": SYNONYMS["answers"]},
        {"topic": "Analogies",      "answers": ANALOGIES["answers"]},
        {"topic": "Palindromes",    "answers": PALINDROMES["answers"]},
        {"topic": "Idioms",         "answers": IDIOMS["answers"]},
        {"topic": "Colloquialisms", "answers": COLLOQUIALISMS["answers"]},
    ]
    answer_key(doc, all_sections)
    scoring_guide(doc)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    out = "/Users/pratappadhy/Documents/Projects/ClaudeCode/.claude/skills/jac/JAC_Grade5_Practice_Worksheet.docx"
    build(out)
