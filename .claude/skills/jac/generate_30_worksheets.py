"""
Generates 30 daily practice worksheets covering ALL 14 topics from the JAC MEGA Test.
Output: worksheets/Day_01_Week1_Foundation.docx … Day_30_Week4_Master.docx
Layout: 30 questions per day matching the actual test format, 4-week progressive difficulty.

Topic distribution per worksheet (mirrors the test's topic spread):
  Science      5   Maths        3   Humanities   3
  Synonyms     3   Antonyms     2   Analogies    2
  Definitions  2   Idioms       2   Colloquials  2
  Palindromes  2   Suffixes     1   Prefixes     1
  Geography    1   Superlatives 1
  ─────────────────────────────────────────────────
  Total       30
"""

import sys, os
sys.path.insert(0, os.path.dirname(__file__))

from question_bank import (
    SCIENCE, MATHEMATICS, HUMANITIES, SYNONYMS, ANTONYMS,
    ANALOGIES, DEFINITIONS, IDIOMS, COLLOQUIALISMS,
    PALINDROMES, SUFFIXES, PREFIXES, GEOGRAPHY, SUPERLATIVES
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Shared helpers ────────────────────────────────────────────────────────────

def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def hide_borders(table):
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:sz"), "0")
        tb.append(el)
    tbl_pr.append(tb)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:tcW")):
            tc_pr.remove(old)
        tcw = OxmlElement("w:tcW")
        tcw.set(qn("w:w"), str(int(width_cm * 567)))
        tcw.set(qn("w:type"), "dxa")
        tc_pr.append(tcw)

def cell_vert_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tc_pr.append(va)

def para_space(para, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"),  str(int(after_pt * 20)))
    pPr.append(spc)

def rn(para, text, bold=False, size=11, colour=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        r, g, b = int(colour[0:2],16), int(colour[2:4],16), int(colour[4:6],16)
        run.font.color.rgb = RGBColor(r,g,b)
    return run

# ── Colours ───────────────────────────────────────────────────────────────────
RED   = "C0392B";  AMBER = "E67E22";  GREEN = "1E8449";  DARK  = "1A5276"
LIGHT = "EAF4FF";  GREY  = "F2F3F4";  WHITE = "FFFFFF";  ALT   = "EBF5FB"
RED_L = "FADBD8";  AMB_L = "FDEBD0";  GRN_L = "D5F5E3"

# Colour per topic (matches priority from original test result)
TOPIC_COLOUR = {
    "Science":        RED,    # was weakest — 5 wrong
    "Synonyms":       RED,    # 2 wrong
    "Analogies":      RED,    # 1 wrong
    "Palindromes":    RED,    # 2 wrong
    "Idioms":         AMBER,  # 1 wrong
    "Colloquialisms": AMBER,  # 1 wrong
    "Mathematics":    GREEN,  # all correct
    "Humanities":     GREEN,
    "Antonyms":       GREEN,
    "Definitions":    GREEN,
    "Suffixes":       GREEN,
    "Prefixes":       GREEN,
    "Geography":      GREEN,
    "Superlatives":   GREEN,
}

WEEK_CONFIG = {
    1: {"name": "Foundation", "colour": RED,
        "note": "No time limit. Focus on understanding each question type before answering."},
    2: {"name": "Building",   "colour": AMBER,
        "note": "Target: finish within 35 minutes. Check all answers before submitting."},
    3: {"name": "Speed",      "colour": GREEN,
        "note": "Target: finish in 30 minutes. Balance speed AND accuracy — aim for 85%+."},
    4: {"name": "Master",     "colour": DARK,
        "note": "Exam conditions. 30 minutes, no breaks, no hints. Review every mistake carefully."},
}

# 30-question distribution matching the test topic spread
DAILY_PLAN = [
    ("Science",        5),
    ("Mathematics",    3),
    ("Humanities",     3),
    ("Synonyms",       3),
    ("Antonyms",       2),
    ("Analogies",      2),
    ("Definitions",    2),
    ("Idioms",         2),
    ("Colloquialisms", 2),
    ("Palindromes",    2),
    ("Suffixes",       1),
    ("Prefixes",       1),
    ("Geography",      1),
    ("Superlatives",   1),
]

POOL = {
    "Science":        SCIENCE,
    "Mathematics":    MATHEMATICS,
    "Humanities":     HUMANITIES,
    "Synonyms":       SYNONYMS,
    "Antonyms":       ANTONYMS,
    "Analogies":      ANALOGIES,
    "Definitions":    DEFINITIONS,
    "Idioms":         IDIOMS,
    "Colloquialisms": COLLOQUIALISMS,
    "Palindromes":    PALINDROMES,
    "Suffixes":       SUFFIXES,
    "Prefixes":       PREFIXES,
    "Geography":      GEOGRAPHY,
    "Superlatives":   SUPERLATIVES,
}

def pick(pool, day, count):
    n = len(pool)
    start = (day * count) % n
    return [pool[(start + i) % n] for i in range(count)]

# ── Document components ───────────────────────────────────────────────────────

def section_header(doc, title, tag, fill):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hide_borders(tbl)
    set_col_width(tbl, 0, 13.5)
    set_col_width(tbl, 1, 3.5)
    lc, rc = tbl.cell(0,0), tbl.cell(0,1)
    shade_cell(lc, fill); shade_cell(rc, fill)
    cell_vert_center(lc); cell_vert_center(rc)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rn(lp, f"  {title}", bold=True, size=11, colour=WHITE)
    para_space(lp, before_pt=4, after_pt=4)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn(rp, f"{tag}  ", bold=True, size=8.5, colour=WHITE)
    para_space(rp, before_pt=4, after_pt=4)
    g = doc.add_paragraph(); para_space(g, after_pt=1)

def mcq_block(doc, questions, start_num):
    for i, q in enumerate(questions):
        num = start_num + i
        tbl = doc.add_table(rows=2, cols=2)
        hide_borders(tbl)
        set_col_width(tbl, 0, 0.9)
        set_col_width(tbl, 1, 16.1)
        bg = GREY if i % 2 == 0 else WHITE
        for row in tbl.rows:
            for c in row.cells:
                shade_cell(c, bg)
        np = tbl.cell(0,0).paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn(np, f"{num}.", bold=True, size=11, colour="1A1A2E")
        para_space(np, before_pt=4, after_pt=1)
        qp = tbl.cell(0,1).paragraphs[0]
        rn(qp, f"  {q['q']}", size=11, colour="1A1A2E")
        para_space(qp, before_pt=4, after_pt=1)
        op = tbl.cell(1,1).paragraphs[0]
        for j, opt in enumerate(q["opts"]):
            rn(op, f"  ({chr(65+j)}) {opt}    ", size=10.5, colour="2C3E50")
        para_space(op, before_pt=1, after_pt=4)
        sp = doc.add_paragraph(); para_space(sp, after_pt=1)

def answer_table(doc, sections):
    doc.add_page_break()
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(hdr, "ANSWER KEY  —  Parent / Teacher Use Only", bold=True, size=14, colour=DARK)
    para_space(hdr, after_pt=6)

    rows_data = []
    for sec in sections:
        for q in sec["qs"]:
            short = q["q"][:52] + ("…" if len(q["q"]) > 52 else "")
            rows_data.append((sec["label"], short, q["ans"], q["expl"]))

    tbl = doc.add_table(rows=1+len(rows_data), cols=4)
    tbl.style = "Table Grid"
    col_widths = [2.6, 6.2, 1.3, 6.9]
    for ci, w in enumerate(col_widths):
        set_col_width(tbl, ci, w)

    for ci, hd in enumerate(["Topic","Question","Ans","Explanation"]):
        cell = tbl.cell(0, ci)
        shade_cell(cell, DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, hd, bold=True, size=9.5, colour=WHITE)
        para_space(p, before_pt=2, after_pt=2)

    for ri, (topic, q_short, ans, expl) in enumerate(rows_data):
        row = tbl.rows[ri+1]
        bg = ALT if ri % 2 == 0 else WHITE
        for cell in row.cells:
            shade_cell(cell, bg)
        data = [topic, q_short, ans, expl]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT]*4
        aligns[2] = WD_ALIGN_PARAGRAPH.CENTER
        for ci, val in enumerate(data):
            p = row.cells[ci].paragraphs[0]
            p.alignment = aligns[ci]
            rn(p, val, bold=(ci==2), size=9,
               colour=RED if ci==2 else "1A1A2E")
            para_space(p, before_pt=2, after_pt=2)

# ── Build one day ─────────────────────────────────────────────────────────────

def build_day(day: int, output_dir: str) -> str:
    week     = min((day // 7) + 1, 4)
    wk       = WEEK_CONFIG[week]
    day_label = day + 1
    date_str  = f"Day {day_label:02d}  —  Week {week}: {wk['name']}"

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21);   sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.2);  sec.right_margin  = Cm(2.2)
    sec.top_margin    = Cm(2.0);  sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── Header bar ───────────────────────────────────────────────────────────
    htbl = doc.add_table(rows=1, cols=2)
    hide_borders(htbl)
    set_col_width(htbl, 0, 12); set_col_width(htbl, 1, 5)
    shade_cell(htbl.cell(0,0), wk["colour"])
    shade_cell(htbl.cell(0,1), wk["colour"])
    lp = htbl.cell(0,0).paragraphs[0]
    rn(lp, f"  GENERAL ABILITY  —  {date_str}", bold=True, size=12, colour=WHITE)
    para_space(lp, before_pt=5, after_pt=5)
    rp = htbl.cell(0,1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn(rp, "James An College  |  Grade 5  ", size=9, colour=WHITE)
    para_space(rp, before_pt=5, after_pt=5)

    # ── Student fields ────────────────────────────────────────────────────────
    doc.add_paragraph()
    ftbl = doc.add_table(rows=2, cols=4)
    ftbl.style = "Table Grid"
    ftbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, lbl in enumerate(["Name","Date","Score  / 30","Time"]):
        shade_cell(ftbl.cell(0,ci), DARK)
        p = ftbl.cell(0,ci).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, lbl, bold=True, size=9.5, colour=WHITE)
        para_space(p, before_pt=2, after_pt=2)
        shade_cell(ftbl.cell(1,ci), LIGHT)
        ftbl.cell(1,ci).paragraphs[0].add_run(" " * 30)

    # ── Week note ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    ntbl = doc.add_table(rows=1, cols=1)
    hide_borders(ntbl)
    shade_cell(ntbl.cell(0,0), LIGHT)
    np2 = ntbl.cell(0,0).paragraphs[0]
    rn(np2, f"Week {week} — {wk['name']}:  ", bold=True, size=10, colour=DARK)
    rn(np2, wk["note"], size=10, italic=True, colour="2C3E50")
    para_space(np2, before_pt=4, after_pt=4)
    doc.add_paragraph()

    # ── Questions ─────────────────────────────────────────────────────────────
    q_num    = 1
    all_secs = []

    for topic, count in DAILY_PLAN:
        qs = pick(POOL[topic], day, count)
        fill = TOPIC_COLOUR[topic]
        tag  = ("Priority ★★★" if fill == RED
                else "Improve ★★" if fill == AMBER
                else "Good ✓")
        section_header(doc, topic, tag, fill)
        mcq_block(doc, qs, q_num)
        q_num += count
        all_secs.append({"label": topic, "qs": qs})

    # ── Answer key ────────────────────────────────────────────────────────────
    answer_table(doc, all_secs)

    fname = os.path.join(output_dir,
            f"Day_{day_label:02d}_Week{week}_{wk['name']}.docx")
    doc.save(fname)
    return fname

# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(__file__), "worksheets")
    os.makedirs(out_dir, exist_ok=True)

    # Remove old worksheets
    for f in os.listdir(out_dir):
        if f.endswith(".docx"):
            os.remove(os.path.join(out_dir, f))

    print(f"Generating 30 worksheets  →  {out_dir}\n")
    for day in range(30):
        path = build_day(day, out_dir)
        print(f"  [{day+1:2d}/30]  {os.path.basename(path)}")

    print(f"\nDone. 30 worksheets saved to:\n  {out_dir}")
